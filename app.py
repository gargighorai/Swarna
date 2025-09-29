# Standard Library imports
import io
import os
import json
from datetime import date
from io import BytesIO
from datetime import datetime
# Third-party imports
import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from flask import Flask, jsonify, render_template, request, redirect, send_file, url_for, flash,Blueprint
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from flask_migrate import Migrate
from flask_sqlalchemy import SQLAlchemy
from itsdangerous import URLSafeTimedSerializer
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from sqlalchemy.orm import joinedload
from docx import Document
from docx.shared import Inches,Pt,Cm,RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_TAB_ALIGNMENT, WD_TAB_LEADER
from models import db, User, Patient, Advice, Drug
#---------------------------------
# App setup
# ---------------------------
# Configure your upload folder
UPLOAD_FOLDER = 'temp_uploads'
ALLOWED_EXTENSIONS = {'json'}
app = Flask(__name__)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
# Essential configurations
basedir = os.path.abspath(os.path.dirname(__file__))
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///'+ os.path.join(basedir, 'site.db')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False  # Set this to False
app.config['SECRET_KEY'] = "new-secret-key" # This is critical
app.config['DEBUG'] = True  # Enable debug mode for development

db.init_app(app)

# Create the upload folder if it doesn't exist
if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# --- Helper Function to Apply Shading (Background Color) ---
def set_table_border_color(table, color="D3D3D3"):  # hex color for light gray
    tbl = table._tbl  # access the XML
    tblPr = tbl.tblPr
     
    tblBorders = OxmlElement('w:tblBorders')
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = tblBorders.find(qn(f'w:{border_name}'))
        if border is None:
            border = OxmlElement(f'w:{border_name}')
            tblBorders.append(border)
        border.set(qn('w:val'), 'single')       # line style
        border.set(qn('w:sz'), '4')             # width (twips)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)        # border color
 
def format_table_header(table, bg="0070C0", text=(255, 255, 255)):
    hdr_cells = table.rows[0].cells
    for cell in hdr_cells:
        # Save original text
        original_text = cell.text.strip()
        cell.text = ""  # clear
        run = cell.paragraphs[0].add_run(original_text)
        run.font.color.rgb = RGBColor(*text)
        # Set background
        set_cell_background(cell, bg)
def set_cell_background(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

# Initialize extensions
migrate = Migrate(app, db)
login_manager = LoginManager()
login_manager.login_view = "login"
login_manager.init_app(app)
# ---------------------------
# Flask-Login
# ---------------------------
@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))
admin_bp = Blueprint('admin', __name__, url_prefix='/admin')
# Assuming your DRUG_FILE is already defined at the top of your file
DRUG_FILE = os.path.join("static", "drugs.json")

def load_drugs_from_static_file():
    imported_count = 0
    try:
        with open(DRUG_FILE, "r") as file:
            drugs_data = json.load(file)
            for drug_entry in drugs_data:
                name = drug_entry.get('name')
                if name:
                    existing = Drug.query.filter_by(name=name).first()
                    if not existing:
                        db.session.add(Drug(name=name))
                        imported_count += 1
            db.session.commit()
    except FileNotFoundError:
        db.session.rollback()
        print(f"Error: The file '{DRUG_FILE}' was not found.")
        return 0
    except Exception as e:
        db.session.rollback()
        print(f"An unexpected error occurred while loading drugs: {e}")
        return 0
    return imported_count

@app.route('/import_static_drugs')
def import_static_drugs():
    imported_count = load_drugs_from_static_file()    
    if imported_count > 0:
        flash(f"Successfully imported {imported_count} new drugs.", 'success')
    else:
        flash("No new drugs were imported. The database may already be up to date.", 'info')
    return redirect(url_for('admin_drugs'))

# This is the route you want to redirect to
@app.route('/admin/drugs')
def admin_drugs():
    # ... your code to fetch drugs and render the template
    return render_template('admin_drugs.html')
# A function to import drugs directly into your database
@app.route('/admin/drugs/import_drugs', methods=['POST'])
@login_required
def import_drugs():
# 1. Handle file upload
    if 'drugs_file' not in request.files:
        flash('No file part', 'danger')
        return redirect(url_for('manage_drugs'))   
    file = request.files['drugs_file']   
    if file.filename == '':
        flash('No selected file', 'warning')
        return redirect(url_for('manage_drugs'))
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
       # 2. Save the uploaded file temporarily
        try:
            file.save(filepath)           
            # 3. Process the file and import to database
            imported_count = import_json_to_db(filepath)           
            if imported_count > 0:
                flash(f'Successfully imported {imported_count} drugs.', 'success')
            else:
                flash('No valid drugs found in the file.', 'warning')            
        except Exception as e:
            flash(f'An error occurred during import: {e}', 'danger')
            return redirect(url_for('manage_drugs'))            
        finally:
            # 4. Clean up the temporary file
            if os.path.exists(filepath):
                os.remove(filepath)        
        return redirect(url_for('manage_drugs'))
    flash('Invalid file type.', 'danger')
    return redirect(url_for('manage_drugs'))
# A helper function to do the database import
def import_json_to_db(filepath):
    imported_count = 0    
    try:
        with open(filepath, 'r') as f:
            drugs_data = json.load(f)
            
        for drug_entry in drugs_data:
            try:
                # Use .get() to handle missing keys gracefully
                name = drug_entry.get('name')
                
                # Ensure the name is not None before proceeding
                if not name:
                    continue
                
                # Check if a drug with the same name already exists
                existing_drug = Drug.query.filter_by(name=name).first()
                if existing_drug:
                    # You could update it here, or just skip it
                    print(f"Skipping duplicate drug: {name}")
                    continue
                
                new_drug = Drug(name=name)
                db.session.add(new_drug)
                imported_count += 1
                
            except Exception as e:
                # Rollback on a per-entry basis
                db.session.rollback()
                print(f"Failed to import a drug entry: {e}")
                
        db.session.commit()
    
    except (json.JSONDecodeError, FileNotFoundError) as e:
        print(f"File processing error: {e}")
        db.session.rollback()
        
    return imported_count

@app.route('/admin/drugs')
@login_required
def manage_drugs():
    drugs = Drug.query.order_by(Drug.name).all()
    return render_template("admin_drugs.html", drugs=drugs)
@app.route("/add_drugs")
@app.route('/admin/drugs/delete/<int:drug_id>', methods=["POST"])
@login_required
def delete_drug(drug_id):
    drug = Drug.query.get_or_404(drug_id)
    db.session.delete(drug)
    db.session.commit()
    flash("Drug deleted", "danger")
    return redirect(url_for("manage_drugs"))
@app.route('/admin/drugs/edit/<int:drug_id>', methods=["POST"])
@login_required
def edit_drug(drug_id):
    drug = Drug.query.get_or_404(drug_id)
    new_name = request.form.get("name", "").strip()
    if new_name:
        if not Drug.query.filter_by(name=new_name).first():
            drug.name = new_name
            db.session.commit()
            flash("Drug updated successfully.", "success")
        else:
            flash("Drug with this name already exists.", "warning")
    else:
        flash("Invalid drug name.", "danger")
    return redirect(url_for("manage_drugs"))
@app.route('/admin/drugs/export')
@login_required
def export_drugs():
    drugs = Drug.query.all()
    data = [{"id": drug.id, "name": drug.name} for drug in drugs]
    response = app.response_class(
        response=json.dumps(data, indent=2),
        mimetype='application/json'
    )
    response.headers.set("Content-Disposition", "attachment", filename="drugs.json")
    return response
# ---------------------------
# Routes for user authentication 
# ---------------------------
@app.route("/")
def home():
    if current_user.is_authenticated:
        return redirect(url_for("dashboard"))
    return redirect(url_for("login"))

@app.route('/')
def index():
    try:
        patients= Patient.query.options(joinedload(Patient.advice).joinedload(Advice.prescribed_drugs)).all()
        flash("Eager loading is working correctly!","success")
        return "your"
    except Exception as e:
        return f"database error: {e}" ,500
    
@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        # Manually get data from the form
        username = request.form['username']
        email = request.form['email']
        password = request.form['password_hash']
        degree = request.form['degree']
        doc_mob = request.form['doc_mob']
        reg_no = request.form['reg_no']
        website = request.form['website']

        # Manual validation (you must add more checks here)
        if not username or not email or not password:
            flash('All fields are required.', 'danger')
            return redirect(url_for('register'))

        # Check for existing users
        user_by_username = User.query.filter_by(username=username).first()
        user_by_email = User.query.filter_by(email=email).first()
        if user_by_username or user_by_email:
            flash('Username or email already exists.', 'danger')
            return redirect(url_for('register'))
        # Hash the password
        hashed_password = generate_password_hash(password)
        # Create and save the new user
        new_user = User(
            username=username,
            email=email,
            password_hash=hashed_password,
            degree=degree,
            doc_mob=doc_mob,
            reg_no=reg_no,
            website=website
        )
        db.session.add(new_user)
        db.session.commit()
        flash('Registration successful! You can now log in.', 'success')
        return redirect(url_for('login'))
    return render_template('register.html')
    
@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        email = request.form["email"]
        password_hash = request.form["password"]
        user = User.query.filter_by(email=email).first()
        if user and check_password_hash(user.password_hash, password_hash):
            login_user(user)
            return redirect(url_for("dashboard"))
        else:
            flash("Invalid credentials", "danger")
    return render_template("login.html")

@app.route("/logout")
@login_required
def logout():
    logout_user()
    return redirect(url_for("login"))

@app.route("/dashboard")
#@login_required
def dashboard():
    patients = Patient.query.filter_by(user_id=current_user.id).all()
    advices=Patient.query.filter_by(user_id=current_user.id).all()   
    return render_template("dashboard.html", patients=patients,advices=advices)
# --------Routes for patient management ------------
@app.route("/add_patient", methods=["GET", "POST"])
@login_required
def add_patient():
    if request.method == "POST":
        name = request.form["name"]
        age = request.form.get("age")
        gender = request.form.get("gender")
        mob_no= request.form.get("mob_no")
        address= request.form.get("address")
        new_patient = Patient(
                          name=name, 
                          age=age,
                          gender=gender,
                          mob_no=mob_no,
                          address=address,
                          user_id=current_user.id)
        db.session.add(new_patient)
        db.session.commit()
        # Redirect the user to the dashboard route
        return redirect(url_for('dashboard'))
    
    return render_template("add_patient.html")
from datetime import date
@app.route('/patient/<int:patient_id>')
def patient_profile(patient_id):
    # Find the patient by ID.
    patient=Patient.query.get_or_404(patient_id)
    advices = Advice.query.filter_by(patient_id=patient.id).order_by(Advice.timestamp.desc()).all()
    today=date.today()   
    return render_template('patient_profile.html', patient=patient, today=today)

@app.route('/print_advice/<int:advice_id>')
def print_advice(advice_id):
    now=date.today()
    patient = Patient.query.options(db.joinedload(Patient.advice).joinedload(Advice.prescribed_drugs)).get_or_404(advice_id)
    advice = Advice.query.options(db.joinedload(Advice.prescribed_drugs)).get_or_404(advice_id)
    return render_template('print_advice.html',patient=patient, advice=advice,today=now)

@app.route('/delete_advice/<int:advice_id>', methods=['POST'])
def delete_advice(advice_id):
    advice_to_delete = Advice.query.get_or_404(advice_id)
    
    patient_id = advice_to_delete.patient_id
    
    # Delete the advice from the database
    db.session.delete(advice_to_delete)
    db.session.commit()
    
    flash("Advice deleted successfully!", "success")
    return redirect(url_for('patient_profile', patient_id=patient_id))
@app.route("/delete_patient/<int:patient_id>",methods=["GET","POST"])
@login_required
def delete_patient(patient_id):
    patient= Patient.query.get_or_404(patient_id)
    db.session.delete(patient)
    db.session.commit()
    flash("Patient data deleted","success")
    return redirect(url_for("dashboard"))

@app.route('/full_advice/<int:patient_id>', methods=['GET', 'POST'])
@login_required
def full_advice(patient_id):
    patient = Patient.query.get_or_404(patient_id)
    drugs = Drug.query.all()

    if request.method == 'POST':
        notes = request.form.get('notes')
        prescribed_drugs = request.form.get('prescribed_drugs')
        prescribed_drugs = json.loads(prescribed_drugs) if prescribed_drugs else []

        # Save advice
        advice = Advice(patient_id=patient.id, doctor_id=current_user.id, notes=notes)
        db.session.add(advice)
        db.session.commit()
        # Save prescribed drugs
        db.session.commit()

        flash("Advice saved successfully!", "success")
        return redirect(url_for('dashboard'))

    return render_template('print_full_advice.html', patient=patient, drugs=drugs)    
from datetime import datetime

#--------Certificate and Receipt Routes ------------
from datetime import datetime
@app.route('/death_certificate/',methods=['GET', 'POST'])
@login_required
def death_certificate():
    return render_template('death_certificate.html',
            now=datetime.now())
@app.route("/certificate/medical")
@login_required
def medical_certificate():
    now=datetime.now()
    return render_template("certificate/medical_certificate.html",
                           now=datetime.now())

@app.route("/certificate/fitness")
@login_required
def fitness_certificate():
    return render_template("certificate/fitness_certificate.html")

@app.route("/certificate/custom")
@login_required
def custom_certificate():
    return render_template("certificate/custom_certificate.html")

@app.route('/receipt')
@login_required
def receipt():
   return  render_template('payment_receipt.html',now=datetime.now())
import os, json
from flask import current_app
# ------------routes for giving advice ------------
from flask import render_template, request, flash, redirect, url_for

@app.route('/give_advice/<int:patient_id>', methods=['GET', 'POST'])
def give_advice(patient_id):
    patient = Patient.query.get_or_404(patient_id) 
    drugs= Drug.query.order_by(Drug.name).all()  
    # This block handles the form submission (POST request)
    if request.method == 'POST':
        selected_drugs= request.form.getlist('drugs')
        #Create advice entry
        advice=Advice(patient_id=patient.id)
        db.session.add(advice)
        db.session.commit()

        for drug_id in selected_drugs:
            drug= Drug.query.get(int(drug_id))
            if drug:
                advice.prescribed_drugs.append(drug)
        
        db.session.commit()
        flash("Advice successfully saved","success")
        return redirect(url_for('patient_profile', patient_id=patient.id)) 
    available_drugs = Drug.query.all()
    return render_template('give_advice.html', patient=patient, available_drugs=available_drugs)
from flask import jsonify
@app.route("/api/drugs")
def get_drugs():
    # Query the database to get all drug entries
    drugs=Drug.query.order_by(Drug.name).all()
    return jsonify([{"id":d.id, "name":d.name } for d in drugs])
    all_drugs = Drug.query.all()
    # This makes them easily convertible to JSON.
    drugs_list = [
        {"name": drug.name}
        for drug in all_drugs
    ]
    
    # Return the list as a JSON response
    return jsonify(drugs_list)
@app.route("/advice/history/<int:patient_id>")
@login_required
def advice_history(patient_id):
    patient = Patient.query.get_or_404(patient_id)
    advices = Advice.query.filter_by(patient_id=patient.id).order_by(Advice.timestamp.desc()).all()
    return render_template("advice_history.html", patient=patient, advices=advices)

# --- DOCX Generation Logic ---
# New route for the patient data entry form
@app.route('/patient_data_entry', methods=['GET'])
def data_entry_form():
    now=date.today()
    available_drugs = Drug.query.all()
    return render_template('patient_form.html', patient_id=1001, available_drugs=available_drugs,today=now)

# New Route to export as  Docx file --

from flask import send_file
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
from datetime import datetime
@app.route('/create_patient_doc/<int:patient_id>')
def create_patient_doc(patient_id):
    patient = Patient.query.get_or_404(patient_id)
    document = Document() 
    section = document.sections[0]
    section.top_margin = Cm(1)       # 1 cm
    section.bottom_margin = Inches(1) # 1 inch
    section.left_margin = Inches(1.7)   # 1 inch
    section.right_margin = Inches(1)  # 1 inch
   
    # --- Utility: faint / invisible table borders ---
    def set_table_borders(table, color="FFFFFF", size="2"):
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top','left','bottom','right','insideH','insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), size)   # thin line
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), color)  # D3D3D3 = faint gray, "FFFFFF" = invisible
            tblBorders.append(border)
        tblPr.append(tblBorders)

    # --- Header: Doctor Info ---
    header_table = document.add_table(rows=1, cols=2)
    row = header_table.rows[0]
    left = row.cells[0]
    right = row.cells[1]
    doctor=current_user
    left.text = f"{doctor.username}\n{doctor.degree}\n{doctor.reg_no}"
    right.text = f"Mob:{doctor.doc_mob}\n{doctor.website}\nEmail us:{doctor.email}\nDate:{datetime.now().strftime('%d-%m-%Y')}"
    set_table_borders(header_table,  color="FFFFFF", size="2")  # invisible borders
    header_table.autofit=False
    header_table.columns[0].width = Inches(5)   # left column wider
    header_table.columns[1].width = Inches(2)   # right column smaller
    document.add_paragraph()  # spacing

    # --- Patient Info Row ---
    pat_table = document.add_table(rows=1, cols=1)
    pat_table.rows[0].cells[0].text = (
        f" Name: {patient.name}  | "
        f"Age: {patient.age} yrs |"
        f" {patient.gender}  | "
        f"Address:{patient.address}  "
    )
    set_table_borders(pat_table, color="D3D3D3", size="2")  # faint border

    document.add_paragraph()  # spacing

    # --- Two Column Layout ---
   
    main_table = document.add_table(rows=1, cols=2)
    main_table.autofit = False
    for row in main_table.rows:
       row.cells[0].width = Inches(2.5) 
       row.cells[1].width=Inches(7.5) # left column (Vitals + Complaints)
       #row.cells[3].width = Inches(7.5)  # right column (Advice)
    #main_table.columns[0].width =Inches(1.5)
    #main_table.columns[1].width =Inches(4.5)
    left_col = main_table.rows[0].cells[0]
    #shading = main_table.cell(0,1)._tc.get_or_add_tcPr().get_or_add_shd()
    #shading.set(qn('w:fill'), "D3D3D3")   # light gray fill
    right_col = main_table.rows[0].cells[1]

    # Left: Patient Vitals + Chief complaints
    vitals_text = f"Vitals & complaints:\nBP: \nPulse: 'N/A'\nTemp: "
    complaints_text = f"\nChief Complaints:\n{ 'N/A'}"
    left_col.text = vitals_text + complaints_text

    # Right: Advice + Prescribed Drugs
    advice_text = ""
    for advice in patient.advice:
        advice_text += f"Advice:\n"
        if advice.prescribed_drugs:
            for drug in advice.prescribed_drugs:
                advice_text += f" - {drug.name}\n"
        advice_text += "\n"
    right_col.text = advice_text.strip()

    set_table_borders(main_table,  color="FFFFFF", size="2")  # faint border
    main_table.autofit=False
    # --- Save to BytesIO for Flask response ---
    file_stream = io.BytesIO()
    filename = f"{patient.name}_advice.docx"
    document.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# ---------------------------
# Run and creates all tables if they don’t exist
if __name__ == "__main__":
 with app.app_context():
    db.create_all()
    app.run(debug=True, port=5050)
